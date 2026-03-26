import re
from pathlib import Path
from typing import Optional, List
from datetime import datetime
import requests
import io

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from src.logger import get_logger

logger = get_logger()

class OutputConverter:
    """输出格式转换器"""
    
    @staticmethod
    def save_markdown(
        markdown_content: str,
        output_path: Path,
        include_metadata: bool = True
    ) -> bool:
        """
        保存为Markdown文件
        
        Args:
            markdown_content: Markdown内容
            output_path: 输出路径
            include_metadata: 是否包含元数据
        
        Returns:
            bool: 是否保存成功
        """
        try:
            logger.info(f"保存Markdown文件: {output_path.name}")
            
            # 添加元数据
            if include_metadata:
                metadata = f"---\n"
                metadata += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                metadata += f"生成工具: GLM-OCR OCR转换工具 v1.0.0\n"
                metadata += f"---\n\n"
                content = metadata + markdown_content
            else:
                content = markdown_content
            
            # 创建输出目录
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 写入文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Markdown文件保存成功: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"Markdown文件保存失败: {e}")
            return False
    
    @staticmethod
    def markdown_to_docx(
        markdown_content: str,
        output_path: Path,
        include_metadata: bool = True
    ) -> bool:
        """
        将Markdown内容转换为DOCX文件
        
        Args:
            markdown_content: Markdown内容
            output_path: 输出路径
            include_metadata: 是否包含元数据
        
        Returns:
            bool: 是否转换成功
        """
        try:
            logger.info(f"转换为DOCX: {output_path.name}")
            
            # 创建文档
            doc = Document()
            
            # 添加元数据
            if include_metadata:
                metadata_para = doc.add_paragraph()
                metadata_para.text = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                metadata_para.runs[0].font.size = Pt(10)
                metadata_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                
                metadata_para2 = doc.add_paragraph()
                metadata_para2.text = "生成工具: GLM-OCR OCR转换工具 v1.0.0"
                metadata_para2.runs[0].font.size = Pt(10)
                metadata_para2.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                
                doc.add_paragraph()  # 添加空行
            
            # 解析Markdown并添加到文档
            lines = markdown_content.split('\n')
            current_list_level = 0
            in_code_block = False
            code_block_content = []
            
            for line in lines:
                # 处理Markdown图片: ![alt](url)
                image_match = re.search(r'!\[[^\]]*\]\((https?://[^)\s]+)\)', line.strip())
                if image_match:
                    image_url = image_match.group(1)
                    try:
                        response = requests.get(image_url, timeout=10)
                        if response.status_code == 200:
                            image_data = io.BytesIO(response.content)
                            para = doc.add_paragraph()
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = para.add_run()
                            run.add_picture(image_data, width=Inches(5.8))
                            continue
                    except Exception as e:
                        logger.warning(f"Markdown图片插入失败: {image_url}, 错误: {e}")

                # 处理HTML图片: <img src="..."> 或 <img src='...'>
                html_img_match = re.search(r'<img[^>]+src=[\"\'](https?://[^\"\']+)[\"\']', line.strip(), re.IGNORECASE)
                if html_img_match:
                    image_url = html_img_match.group(1)
                    try:
                        response = requests.get(image_url, timeout=10)
                        if response.status_code == 200:
                            image_data = io.BytesIO(response.content)
                            para = doc.add_paragraph()
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            run = para.add_run()
                            run.add_picture(image_data, width=Inches(5.8))
                            continue
                    except Exception as e:
                        logger.warning(f"HTML图片插入失败: {image_url}, 错误: {e}")

                # 处理代码块
                if line.strip().startswith('```'):
                    if in_code_block:
                        # 代码块结束
                        code_para = doc.add_paragraph()
                        code_para.style = 'Normal'
                        for code_line in code_block_content:
                            code_para.add_run(code_line + '\n')
                        code_block_content = []
                        in_code_block = False
                    else:
                        # 代码块开始
                        in_code_block = True
                    continue
                
                if in_code_block:
                    code_block_content.append(line)
                    continue
                
                # 处理标题
                if line.startswith('# '):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    para = doc.add_heading(title, level=min(level, 9))
                    continue
                
                if line.startswith('## '):
                    para = doc.add_heading(line[3:].strip(), level=2)
                    continue
                
                if line.startswith('### '):
                    para = doc.add_heading(line[4:].strip(), level=3)
                    continue
                
                if line.startswith('#### '):
                    para = doc.add_heading(line[5:].strip(), level=4)
                    continue
                
                # 处理列表
                if line.startswith('- ') or line.startswith('* '):
                    para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                    continue
                
                if line.startswith('  - ') or line.startswith('  * '):
                    para = doc.add_paragraph(line[4:].strip(), style='List Bullet 2')
                    continue
                
                # 处理数字列表
                match = re.match(r'^(\d+)\.\s+(.*)', line)
                if match:
                    para = doc.add_paragraph(match.group(2).strip(), style='List Number')
                    continue
                
                # 处理表格（简单处理，识别管道符）
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if cells:
                        # 检查是否是表格分隔行
                        if all(c.replace('-', '').replace(':', '').replace(' ', '') == '' for c in cells):
                            continue
                        
                        # 添加表格行
                        cols = len(cells)
                        if not hasattr(OutputConverter, '_current_table') or OutputConverter._current_table is None:
                            OutputConverter._current_table = doc.add_table(rows=1, cols=cols)
                            OutputConverter._current_table.style = 'Light Grid Accent 1'
                        
                        row = OutputConverter._current_table.add_row()
                        for i, cell in enumerate(cells):
                            row.cells[i].text = cell
                        continue
                
                OutputConverter._current_table = None
                
                # 处理普通段落
                if line.strip():
                    para = doc.add_paragraph(line)
                else:
                    # 空行
                    if doc.paragraphs and doc.paragraphs[-1].text:
                        doc.add_paragraph()
            
            # 创建输出目录
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"DOCX文件保存成功: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"DOCX转换失败: {e}")
            return False
    
    @staticmethod
    def add_images_to_docx(
        doc_path: Path,
        image_urls: List[str],
        output_path: Optional[Path] = None
    ) -> bool:
        """
        将图片添加到DOCX文档
        
        Args:
            doc_path: 源DOCX文件路径
            image_urls: 图片URL列表
            output_path: 输出路径（如果为None则覆盖原文件）
        
        Returns:
            bool: 是否添加成功
        """
        if output_path is None:
            output_path = doc_path
        
        try:
            logger.info(f"向DOCX添加 {len(image_urls)} 张图片")
            
            doc = Document(doc_path)
            
            for image_url in image_urls:
                try:
                    # 下载图片
                    response = requests.get(image_url, timeout=10)
                    if response.status_code == 200:
                        image_data = io.BytesIO(response.content)
                        # 添加图片到文档
                        para = doc.add_paragraph()
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = para.add_run()
                        run.add_picture(image_data, width=Inches(6))
                        logger.info(f"已添加图片: {image_url}")
                except Exception as e:
                    logger.warning(f"添加图片失败: {image_url}, 错误: {e}")
                    continue
            
            doc.save(output_path)
            logger.info(f"图片添加完成: {output_path}")
            return True
        
        except Exception as e:
            logger.error(f"添加图片到DOCX失败: {e}")
            return False
    
    @staticmethod
    def export_result(
        markdown_content: str,
        output_filename: str,
        output_dir: Path,
        export_format: str = "markdown",
        image_urls: Optional[List[str]] = None,
        include_metadata: bool = True
    ) -> tuple:
        """
        导出OCR识别结果
        
        Args:
            markdown_content: Markdown内容
            output_filename: 输出文件名（不含扩展名）
            output_dir: 输出目录
            export_format: 导出格式 ("markdown"/"docx"/"both")
            image_urls: 图片URL列表（用于DOCX）
            include_metadata: 是否包含元数据
        
        Returns:
            tuple: (成功状态, 输出文件路径列表)
        """
        output_files = []
        
        try:
            output_dir.mkdir(parents=True, exist_ok=True)
            
            if export_format in ("markdown", "both"):
                # 导出为Markdown
                md_path = output_dir / f"{output_filename}.md"
                if OutputConverter.save_markdown(markdown_content, md_path, include_metadata):
                    output_files.append(str(md_path))
            
            if export_format in ("docx", "both"):
                # 导出为DOCX
                docx_path = output_dir / f"{output_filename}.docx"
                if OutputConverter.markdown_to_docx(markdown_content, docx_path, include_metadata):
                    output_files.append(str(docx_path))
            
            logger.info(f"导出完成，输出文件数: {len(output_files)}")
            return True, output_files
        
        except Exception as e:
            logger.error(f"导出失败: {e}")
            return False, []
